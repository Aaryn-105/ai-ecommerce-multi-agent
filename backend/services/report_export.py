"""Report export service — generates PDF (ReportLab) and DOCX (python-docx) files.

Usage::
    svc = ReportExportService()
    pdf_bytes = svc.to_pdf("Report Title", "Summary text", sections_dict)
    docx_bytes = svc.to_docx("Report Title", "Summary text", sections_dict)
"""

from __future__ import annotations

import html

import io

import os

import re

import tempfile

from typing import Any

from backend.core.config import settings

from pathlib import Path

from reportlab.pdfbase import pdfmetrics

from reportlab.pdfbase.ttfonts import TTFont

# ── Optional dependency helpers ──────────────────────────

# CJK font registration - PDF export needs a Chinese-capable font.

_CJK_FONT_NAME = 'AppCJK'

_CJK_BOLD_FONT_NAME = 'AppCJK-Bold'

_FONTS_REGISTERED = False

def _register_cjk_fonts():
    'Register Microsoft YaHei or fallback CJK font. Idempotent.'
    global _FONTS_REGISTERED
    if _FONTS_REGISTERED:
        return
    fonts_dir = Path('C:/Windows/Fonts')
    candidates = [
        ('AppCJK', 'msyh.ttc', 0),
        ('AppCJK-Bold', 'msyhbd.ttc', 0),
        ('AppCJK', 'simsun.ttc', 0),
        ('AppCJK', 'simhei.ttf', None),
    ]
    for name, filename, sub_index in candidates:
        font_path = fonts_dir / filename
        if not font_path.exists():
            continue
        try:
            kwargs = {}
            if sub_index is not None:
                kwargs['subfontIndex'] = sub_index
            pdfmetrics.registerFont(TTFont(name, str(font_path), **kwargs))
        except Exception:
            continue
    _FONTS_REGISTERED = True
_CJK_TABLE_FONT = 'Helvetica'

_CJK_TABLE_BOLD = 'Helvetica-Bold'

def _resolve_cjk_table_font():
    global _CJK_TABLE_FONT, _CJK_TABLE_BOLD
    _register_cjk_fonts()
    registered = pdfmetrics.getRegisteredFontNames()
    if _CJK_FONT_NAME in registered:
        _CJK_TABLE_FONT = _CJK_FONT_NAME
    if _CJK_BOLD_FONT_NAME in registered:
        _CJK_TABLE_BOLD = _CJK_BOLD_FONT_NAME
    return _CJK_TABLE_FONT
def _check_reportlab() -> None:
    try:
        import reportlab  # noqa: F401
    except ImportError:
        msg = "reportlab is required for PDF export. Install: pip install reportlab"
        raise ImportError(msg)
def _check_docx() -> None:
    try:
        import docx  # noqa: F401
    except ImportError:
        msg = "python-docx is required for DOCX export. Install: pip install python-docx"
        raise ImportError(msg)
# ── Section label mapping (Chinese) ──────────────────────

_SECTION_LABELS: dict[str, str] = {
    "product_analysis": "商品选品分析",
    "trend_forecast": "销售趋势预测",
    "competitor_analysis": "竞品对比分析",
    "marketing_copy": "营销文案",
    "inventory": "库存补货建议",
    "pricing": "定价建议",
    "promotion": "促销方案",
}

_AGENT_ORDER = [
    "product_analysis",
    "trend_forecast",
    "competitor_analysis",
    "marketing_copy",
    "inventory",
    "pricing",
    "promotion",
]

def _fmt(val: Any, decimals: int = 2) -> str:
    """Format a number nicely, or return the value as-is."""
    if isinstance(val, float):
        return f"{val:.{decimals}f}"
    if isinstance(val, int):
        return str(val)
    return str(val)
# ═══════════════════════════════════════════════════════════

#  PDF Generation (ReportLab)

# ═══════════════════════════════════════════════════════════

def _markdown_inline(text: str) -> str:
    """Convert report Markdown emphasis to safe ReportLab markup."""
    value = html.escape(str(text), quote=False)
    value = re.sub(r"`([^`]+)`", r'<font color="#374151">\1</font>', value)
    value = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", value)
    value = re.sub(r"__(.+?)__", r"<b>\1</b>", value)
    return re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"<i>\1</i>", value)
_CHART_HEADING_KEYWORDS: dict[str, tuple[str, ...]] = {
    "product_analysis": (
        "综合评分结果",
        "综合评分三级标签",
        "综合评分",
        "评分排名",
    ),
    "trend_forecast": (
        "趋势预测与多维交叉分析",
        "趋势预测分析",
        "销售趋势预测",
        "趋势预测",
        "销售预测",
    ),
    "competitor_analysis": (
        "竞品对比",
        "竞品分析",
        "竞争对手",
    ),
    "marketing_copy": (
        "营销文案",
        "营销策略",
        "文案",
    ),
    "inventory": (
        "库存策略",
        "库存建议",
        "补货建议",
        "库存健康",
    ),
    "pricing": (
        "定价策略",
        "定价建议",
        "价格策略",
        "定价对比",
    ),
    "promotion": (
        "促销建议",
        "促销策略",
        "促销方案",
    ),
}

def _chart_heading_matches(agent_key: str, heading_text: str) -> bool:
    """Return whether a report heading is the natural home for an agent chart."""
    normalized_heading = re.sub(r"\s+", "", str(heading_text)).lower()
    return any(
        keyword.lower() in normalized_heading
        for keyword in _CHART_HEADING_KEYWORDS.get(agent_key, ())
    )
_CJK_TTF: str | None = None

_CJK_BOLD_TTF: str | None = None

_CJK_FONT_TRIED = False

def _load_cjk_fonts():
    """Lazy-load CJK-capable fonts via Pillow. Idempotent."""
    global _CJK_TTF, _CJK_BOLD_TTF, _CJK_FONT_TRIED
    if _CJK_FONT_TRIED:
        return
    _CJK_FONT_TRIED = True
    try:
        from PIL import ImageFont
        fonts_dir = r"C:\Windows\Fonts"
        if os.path.exists(os.path.join(fonts_dir, "msyh.ttc")):
            _CJK_TTF = fonts_dir + "\\msyh.ttc"
        elif os.path.exists(os.path.join(fonts_dir, "simhei.ttf")):
            _CJK_TTF = fonts_dir + "\\simhei.ttf"
        elif os.path.exists(os.path.join(fonts_dir, "simsun.ttc")):
            _CJK_TTF = fonts_dir + "\\simsun.ttc"
        if os.path.exists(os.path.join(fonts_dir, "msyhbd.ttc")):
            _CJK_BOLD_TTF = fonts_dir + "\\msyhbd.ttc"
        else:
            _CJK_BOLD_TTF = _CJK_TTF
    except Exception:
        pass
def _truncate_label(s, n=10):
    s = str(s or "")
    return s if len(s) <= n else s[: n - 1] + "…"
def _color_hex(c):
    """Accept HexColor or hex string, return "#rrggbb"."""
    try:
        return c.hexval()
    except Exception:
        return str(c) if str(c).startswith("#") else "#" + str(c)
def _wrap_pil_image(pil_img, width_pt=480, dpi=120):
    """Wrap a PIL Image as reportlab.platypus.Image (PNG bytes)."""
    from reportlab.platypus import Image as RLImage
    import io as _io
    if pil_img is None:
        return None
    buf = _io.BytesIO()
    pil_img.save(buf, format="PNG")
    buf.seek(0)
    aspect = pil_img.height / pil_img.width
    h_pt = width_pt * aspect
    return RLImage(buf, width=width_pt, height=h_pt)
def _make_bar_chart(
    title: str,
    labels,
    values,
    *,
    color_hex: str = "#0f3460",
    value_format: str = "{:.1f}",
    width_px: int = 960,
    height_px: int = 360,
):
    """Draw a single-series bar chart using Pillow. Returns PIL.Image."""
    from PIL import Image, ImageDraw, ImageFont
    _load_cjk_fonts()
    try:
        font_title = ImageFont.truetype(_CJK_TTF, 22)
        font_axis = ImageFont.truetype(_CJK_TTF, 16)
        font_label = ImageFont.truetype(_CJK_TTF, 14)
        font_val = ImageFont.truetype(_CJK_TTF, 13)
    except Exception:
        font_title = font_axis = font_label = font_val = ImageFont.load_default()
    safe_labels = [_truncate_label(l, 10) for l in labels]
    safe_values = [float(v or 0) for v in values]
    img = Image.new("RGB", (width_px, height_px), "#fafafa")
    d = ImageDraw.Draw(img)
    margin_left, margin_right = 60, 24
    margin_top, margin_bottom = 60, 70
    plot_w = width_px - margin_left - margin_right
    plot_h = height_px - margin_top - margin_bottom
    d.text((width_px / 2, 18), str(title), fill="#1a1a2e", font=font_title, anchor="mm")
    d.rectangle([margin_left, margin_top, margin_left + plot_w, margin_top + plot_h],
                outline="#cbd5e1", width=1)
    vmin = min(safe_values) if safe_values else 0
    vmax = max(safe_values) if safe_values else 0
    if vmax == vmin:
        vmax = vmin + 1
    # scale so vmax is at top, vmin at bottom (or 0 if all non-negative)
    if vmin >= 0:
        top_val = vmax * 1.18 if vmax > 0 else 1
        bot_val = 0
    else:
        span = vmax - vmin
        top_val = vmax + span * 0.18
        bot_val = vmin
    span = top_val - bot_val
    n = max(len(safe_values), 1)
    def val_to_y(v):
        return margin_top + plot_h - (v - bot_val) / span * plot_h
    zero_y = val_to_y(0) if vmin < 0 else (margin_top + plot_h)
    # 5 gridlines spanning the visible value range
    for i in range(0, 5):
        y = margin_top + plot_h - (plot_h * i / 4)
        val = bot_val + span * i / 4
        d.line([(margin_left, y), (margin_left + plot_w, y)], fill="#e2e8f0", width=1)
        label = value_format.format(val)
        d.text((margin_left - 8, y), label, fill="#475569", font=font_axis, anchor="rm")
    # zero baseline
    if vmin < 0:
        d.line([(margin_left, zero_y), (margin_left + plot_w, zero_y)],
               fill="#94a3b8", width=1)
    slot_w = plot_w / n
    bar_w = min(slot_w * 0.6, 70)
    for i, v in enumerate(safe_values):
        x_center = margin_left + slot_w * (i + 0.5)
        x0 = x_center - bar_w / 2
        x1 = x_center + bar_w / 2
        bar_top = val_to_y(v)
        bar_bot = zero_y
        y_top, y_bot = min(bar_top, bar_bot), max(bar_top, bar_bot)
        d.rectangle([x0, y_top, x1, y_bot], fill=color_hex, outline=color_hex)
        label_y = y_top - 6 if v >= 0 else y_bot + 6
        anchor = "mb" if v >= 0 else "mt"
        d.text((x_center, label_y), value_format.format(v),
               fill="#1a1a2e", font=font_val, anchor=anchor)
        d.text((x_center, margin_top + plot_h + 8), safe_labels[i],
               fill="#475569", font=font_label, anchor="mt")
    d.text((margin_left + plot_w / 2, height_px - 12), "商品",
           fill="#1a1a2e", font=font_axis, anchor="mm")
    return img
def _make_grouped_bar_chart(
    title: str,
    labels,
    series_a,
    series_b,
    *,
    legend_a: str = "当前",
    legend_b: str = "建议",
    color_a_hex: str = "#94a3b8",
    color_b_hex: str = "#e94560",
    value_format: str = "${:.0f}",
    width_px: int = 960,
    height_px: int = 400,
):
    """Draw a two-series grouped bar chart with legend using Pillow."""
    from PIL import Image, ImageDraw, ImageFont
    _load_cjk_fonts()
    try:
        font_title = ImageFont.truetype(_CJK_TTF, 22)
        font_axis = ImageFont.truetype(_CJK_TTF, 16)
        font_label = ImageFont.truetype(_CJK_TTF, 14)
        font_val = ImageFont.truetype(_CJK_TTF, 13)
        font_legend = ImageFont.truetype(_CJK_TTF, 16)
    except Exception:
        font_title = font_axis = font_label = font_val = font_legend = ImageFont.load_default()
    safe_labels = [_truncate_label(l, 8) for l in labels]
    sa = [float(v or 0) for v in series_a]
    sb = [float(v or 0) for v in series_b]
    img = Image.new("RGB", (width_px, height_px), "#fafafa")
    d = ImageDraw.Draw(img)
    margin_left, margin_right = 70, 24
    margin_top, margin_bottom = 70, 70
    plot_w = width_px - margin_left - margin_right
    plot_h = height_px - margin_top - margin_bottom
    d.text((width_px / 2, 18), str(title), fill="#1a1a2e", font=font_title, anchor="mm")
    leg_x, leg_y = width_px - 240, 30
    d.rectangle([leg_x, leg_y, leg_x + 18, leg_y + 14], fill=color_a_hex)
    d.text((leg_x + 24, leg_y + 7), legend_a, fill="#1a1a2e", font=font_legend, anchor="lm")
    d.rectangle([leg_x + 120, leg_y, leg_x + 138, leg_y + 14], fill=color_b_hex)
    d.text((leg_x + 144, leg_y + 7), legend_b, fill="#1a1a2e", font=font_legend, anchor="lm")
    d.rectangle([margin_left, margin_top, margin_left + plot_w, margin_top + plot_h],
                outline="#cbd5e1", width=1)
    vmax = max(sa + sb) if (sa or sb) else 0
    if vmax <= 0:
        vmax = 1
    vmax = vmax * 1.18
    n = max(len(sa), 1)
    for i in range(0, 5):
        y = margin_top + plot_h - (plot_h * i / 4)
        d.line([(margin_left, y), (margin_left + plot_w, y)], fill="#e2e8f0", width=1)
        d.text((margin_left - 8, y), value_format.format(vmax * i / 4),
               fill="#475569", font=font_axis, anchor="rm")
    group_w = plot_w / n
    bar_w = min(group_w * 0.35, 40)
    gap = 4
    for i in range(n):
        x_center = margin_left + group_w * (i + 0.5)
        for offset, (v, hex_color) in enumerate([(sa[i], color_a_hex), (sb[i], color_b_hex)]):
            x0 = x_center - bar_w + offset * (bar_w + gap)
            x1 = x0 + bar_w
            bar_h = (v / vmax) * plot_h if vmax > 0 else 0
            y0 = margin_top + plot_h - bar_h
            y1 = margin_top + plot_h
            d.rectangle([x0, y0, x1, y1], fill=hex_color, outline=hex_color)
        d.text((x_center, margin_top + plot_h + 8), safe_labels[i],
               fill="#475569", font=font_label, anchor="mt")
    d.text((margin_left + plot_w / 2, height_px - 12), "商品",
           fill="#1a1a2e", font=font_axis, anchor="mm")
    return img
def _make_pie_chart(
    title: str,
    labels,
    values,
    *,
    width_px: int = 640,
    height_px: int = 360,
):
    """Draw a pie chart with a side legend using Pillow."""
    from PIL import Image, ImageDraw, ImageFont
    import math
    _load_cjk_fonts()
    try:
        font_title = ImageFont.truetype(_CJK_TTF, 22)
        font_label = ImageFont.truetype(_CJK_TTF, 16)
        font_legend = ImageFont.truetype(_CJK_TTF, 15)
    except Exception:
        font_title = font_label = font_legend = ImageFont.load_default()
    safe_labels = [str(l) for l in labels]
    safe_values = [max(float(v or 0), 0.001) for v in values]
    total = sum(safe_values) or 1
    palette = ["#0f3460", "#e94560", "#22c55e", "#f59e0b", "#94a3b8", "#6366f1"]
    img = Image.new("RGB", (width_px, height_px), "#fafafa")
    d = ImageDraw.Draw(img)
    d.text((width_px / 2, 18), str(title), fill="#1a1a2e", font=font_title, anchor="mm")
    cx, cy = width_px // 3, height_px // 2 + 16
    r = min(width_px // 4, height_px // 2 - 40)
    start = -math.pi / 2
    for i, v in enumerate(safe_values):
        ang = (v / total) * 2 * math.pi
        end = start + ang
        d.pieslice([cx - r, cy - r, cx + r, cy + r],
                   start=math.degrees(start), end=math.degrees(end),
                   fill=palette[i % len(palette)], outline="#ffffff", width=2)
        start = end
    leg_x = cx + r + 40
    leg_y = cy - r
    for i, (lab, val) in enumerate(zip(safe_labels, safe_values)):
        ly = leg_y + 26 * i
        d.rectangle([leg_x, ly, leg_x + 16, ly + 16],
                    fill=palette[i % len(palette)],
                    outline=palette[i % len(palette)])
        pct = val / total * 100
        d.text((leg_x + 22, ly + 8),
               f"{lab}  {val:.0f} ({pct:.0f}%)",
               fill="#1a1a2e", font=font_legend, anchor="lm")
    return img
def _render_pdf_charts(
    elements: list,
    agent_key: str,
    data: dict,
    body_style,
) -> None:
    """Insert Pillow-rendered PNG bar/pie/grouped charts into the PDF for quantitative agents."""
    from reportlab.platypus import Paragraph, Spacer
    PRODUCT_PRIMARY = "#0f3460"
    ACCENT = "#e94560"
    SUCCESS = "#22c55e"
    WARNING = "#f59e0b"
    GREY = "#94a3b8"
    if agent_key == "product_analysis":
        products = data.get("selected_products", []) or []
        if len(products) >= 2:
            items = products[:10]
            labels = [p.get("title", "") for p in items]
            values = [float(p.get("final_score", 0) or p.get("composite_score", 0) or 0) for p in items]
            if any(v > 0 for v in values):
                elements.append(Paragraph("<b> Top 10 商品综合评分对比</b>", body_style))
                elements.append(Spacer(1, 2))
                elements.append(_wrap_pil_image(
                    _make_bar_chart("综合评分（越高越优）", labels, values,
                                    color_hex=PRODUCT_PRIMARY, value_format="{:.1f}"),
                    width_pt=460,
                ))
                elements.append(Spacer(1, 8))
            cat_dist = (data.get("statistics", {}) or {}).get("category_distribution", {}) or {}
            if len(cat_dist) >= 2:
                elements.append(Paragraph("<b> 类目商品数分布</b>", body_style))
                elements.append(Spacer(1, 2))
                elements.append(_wrap_pil_image(
                    _make_pie_chart("类目商品数", list(cat_dist.keys()), list(cat_dist.values())),
                    width_pt=420,
                ))
                elements.append(Spacer(1, 8))
    elif agent_key == "trend_forecast":
        forecasts = data.get("product_forecasts", []) or []
        if len(forecasts) >= 2:
            items = forecasts[:10]
            labels = [fc.get("title", "") for fc in items]
            def _sum30(fc):
                f30 = fc.get("forecast_30d") or fc.get("predicted_30d") or []
                if isinstance(f30, list):
                    return float(sum(f30))
                return float(f30 or 0)
            values_30 = [_sum30(fc) for fc in items]
            values_growth = [float(fc.get("trend_rate", 0) or fc.get("growth_rate", 0) or 0) * 100 for fc in items]
            if any(v > 0 for v in values_30):
                elements.append(Paragraph("<b> 30 天销量预测对比</b>", body_style))
                elements.append(Spacer(1, 2))
                elements.append(_wrap_pil_image(
                    _make_bar_chart("30天预测销量（件）", labels, values_30,
                                    color_hex=ACCENT, value_format="{:.0f}"),
                    width_pt=460,
                ))
                elements.append(Spacer(1, 8))
            if any(v != 0 for v in values_growth):
                color = SUCCESS if any(v >= 0 for v in values_growth) else ACCENT
                elements.append(Paragraph("<b> 增长率对比 (%)</b>", body_style))
                elements.append(Spacer(1, 2))
                elements.append(_wrap_pil_image(
                    _make_bar_chart("增长率 (%)", labels, values_growth,
                                    color_hex=color, value_format="{:.1f}"),
                    width_pt=460,
                ))
                elements.append(Spacer(1, 8))
    elif agent_key == "pricing":
        results = data.get("pricing_results", []) or []
        if len(results) >= 2:
            items = results[:8]
            labels = [r.get("title", "") for r in items]
            cur = [float(r.get("current_price", 0) or 0) for r in items]
            sug = [float(r.get("suggested_price", 0) or 0) for r in items]
            elements.append(Paragraph("<b> 定价对比：当前价 vs 建议价</b>", body_style))
            elements.append(Spacer(1, 2))
            elements.append(_wrap_pil_image(
                _make_grouped_bar_chart(
                    "定价对比（USD）", labels, cur, sug,
                    legend_a="当前价", legend_b="建议价",
                    color_a_hex=GREY, color_b_hex=ACCENT,
                    value_format="${:.0f}",
                ),
                width_pt=460,
            ))
            elements.append(Spacer(1, 8))
    elif agent_key == "inventory":
        plans = data.get("replenishment_plans", []) or []
        if len(plans) >= 2:
            items = plans[:8]
            labels = [pl.get("title", "") for pl in items]
            values = [float(pl.get("suggested_reorder_qty", 0) or 0) for pl in items]
            if any(v > 0 for v in values):
                elements.append(Paragraph("<b> 建议补货数量</b>", body_style))
                elements.append(Spacer(1, 2))
                elements.append(_wrap_pil_image(
                    _make_bar_chart("建议补货量（件）", labels, values,
                                    color_hex=WARNING, value_format="{:.0f}"),
                    width_pt=460,
                ))
                elements.append(Spacer(1, 8))
def _render_product_table(
    elements: list[Any],
    products: list[dict[str, Any]],
    body_style: Any,
    small_style: Any,
) -> None:
    """Render selected products as a table."""
    _resolve_cjk_table_font()
    cjk_font = _CJK_TABLE_FONT
    cjk_bold = _CJK_TABLE_BOLD
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
    if not products:
        return
    def _tier_for(score, top_score):
        if top_score <= 0:
            return "备选"
        ratio = score / top_score
        if ratio >= 0.85:
            return " 强势推荐"
        if ratio >= 0.55:
            return "推荐"
        return "备选"
    top_score = products[0].get("final_score", 0) if products else 0
    header = ["#", "推荐等级", "商品名称", "品类", "价格", "评分(评价数)", "综合得分"]
    rows = [header]
    for i, p in enumerate(products[:10], 1):
        rating_obj = p.get("original_rating", {}) or {}
        rate = rating_obj.get("rate", 0) if isinstance(rating_obj, dict) else 0
        count = rating_obj.get("count", 0) if isinstance(rating_obj, dict) else 0
        rows.append([
            str(i),
            _tier_for(p.get("final_score", 0), top_score),
            p.get("title", "")[:38],
            p.get("category", "")[:12],
            f"${_fmt(p.get('price', 0))}",
            f"{_fmt(rate)} ({count})",
            _fmt(p.get("final_score", 0), 2),
        ])
    t = Table(rows, colWidths=[22, 65, 175, 65, 50, 80, 55])
    t.setStyle(TableStyle([
            ("FONT", (0, 0), (-1, -1), cjk_font),
        ("FONTNAME", (0, 0), (-1, 0), cjk_bold),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("ALIGN", (0, 0), (0, -1), "CENTER"),
        ("ALIGN", (3, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f5f5")]),
    ]))
    elements.append(t)
# ═══════════════════════════════════════════════════════════

#  DOCX Generation (python-docx)

# ═══════════════════════════════════════════════════════════

def _build_docx(
    title: str,
    summary: str,
    sections: dict[str, Any],
    *,
    report_document: dict[str, Any] | None = None,
    content_md: str = "",
) -> bytes:
    """Generate a designed DOCX report from the semantic document model.
    Mirrors ``_build_professional_pdf`` so PDF and DOCX share heading
    ordering, block coverage, and chart placement logic.
    """
    _check_docx()
    from docx import Document
    from backend.services.report_document import (
        build_report_document,
        normalize_report_document,
    )
    if report_document:
        document = normalize_report_document(
            report_document, title=title, summary=summary
        )
    else:
        document = build_report_document(
            content_md, title=title, summary=summary
        )
    doc = Document()
    _setup_docx_page(doc)
    # Cover page (title + summary + meta)
    _render_docx_cover(doc, document, title, summary)
    doc.add_page_break()
    # Flatten _agent_sections so chart collection finds data
    flat_sections = dict(sections or {})
    nested = flat_sections.pop("_agent_sections", None)
    if isinstance(nested, dict):
        for k, v in nested.items():
            flat_sections.setdefault(k, v)
    # Pre-compute charts per agent (mirrors _build_professional_pdf)
    chart_insertions: dict[str, list[tuple[str, Any]]] = {}
    for agent_key in _AGENT_ORDER:
        data = flat_sections.get(agent_key)
        if not isinstance(data, dict):
            continue
        charts = _collect_docx_charts(agent_key, data)
        if charts:
            chart_insertions[agent_key] = charts
    _render_semantic_docx(doc, document, chart_insertions=chart_insertions)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
def _professional_table(
    headers: list[str],
    rows: list[list[str]],
    *,
    table_width: float,
    body_style: Any,
    header_style: Any,
) -> Any:
    """Create a polished, wrapping table from semantic report data."""
    from reportlab.lib import colors
    from reportlab.platypus import LongTable, Paragraph, TableStyle
    column_count = max(len(headers), max((len(row) for row in rows), default=0))
    normalized_headers = headers + [""] * (column_count - len(headers))
    normalized_rows = [row + [""] * (column_count - len(row)) for row in rows]
    lengths = [
        max(
            len(str(normalized_headers[index])),
            max((len(str(row[index])) for row in normalized_rows), default=0),
        )
        for index in range(column_count)
    ]
    weights = [min(max(length, 6), 30) for length in lengths]
    total = sum(weights) or column_count
    widths = [table_width * weight / total for weight in weights]
    table_data = [
        [Paragraph(_markdown_inline(cell), header_style) for cell in normalized_headers],
        *[
            [Paragraph(_markdown_inline(cell), body_style) for cell in row]
            for row in normalized_rows
        ],
    ]
    table = LongTable(table_data, colWidths=widths, repeatRows=1, hAlign="LEFT")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0f3460")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f4f7fb")]),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#cbd5e1")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    return table
def _is_appendix_heading(value: str) -> bool:
    normalized = re.sub(r"\s+", "", str(value or ""))
    return "\u9644\u5f55" in normalized or "appendix" in normalized.lower()
def _render_semantic_report(
    elements: list[Any],
    report_document: dict[str, Any],
    *,
    styles: dict[str, Any],
    table_width: float,
    chart_insertions: dict[str, list[Any]],
) -> None:
    """Render semantic report blocks without invoking the Markdown PDF renderer."""
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, Spacer
    def append_matching_charts(heading_text: str) -> None:
        matched_agents = [
            agent_key
            for agent_key in list(chart_insertions)
            if _chart_heading_matches(agent_key, heading_text)
        ]
        for agent_key in matched_agents:
            elements.extend(chart_insertions.pop(agent_key))
    def append_chart_fallback() -> None:
        if not chart_insertions:
            return
        elements.append(Paragraph("\u6570\u636e\u53ef\u89c6\u5316", styles["section"]))
        for agent_key, chart_elements in list(chart_insertions.items()):
            elements.append(Paragraph(_SECTION_LABELS.get(agent_key, agent_key), styles["subheading"]))
            elements.extend(chart_elements)
            chart_insertions.pop(agent_key, None)
    for section_index, section in enumerate(report_document.get("sections") or []):
        if not isinstance(section, dict):
            continue
        section_title = str(section.get("title") or "\u5206\u6790\u7ed3\u679c").strip()
        if _is_appendix_heading(section_title):
            append_chart_fallback()
        elements.append(Paragraph(_markdown_inline(section_title), styles["section"]))
        append_matching_charts(section_title)
        first_narrative = True
        for block in section.get("blocks") or []:
            if not isinstance(block, dict):
                continue
            block_type = str(block.get("type") or "paragraph")
            if block_type == "heading":
                heading_text = str(block.get("text") or "").strip()
                if not heading_text:
                    continue
                elements.append(Paragraph(_markdown_inline(heading_text), styles["subheading"]))
                append_matching_charts(heading_text)
                first_narrative = False
            elif block_type == "paragraph":
                text = str(block.get("text") or "").strip()
                if not text:
                    continue
                style = styles["lead"] if first_narrative else styles["body"]
                elements.append(Paragraph(_markdown_inline(text), style))
                first_narrative = False
            elif block_type == "callout":
                text = str(block.get("text") or "").strip()
                if text:
                    elements.append(Paragraph(_markdown_inline(text), styles["callout"]))
                    first_narrative = False
            elif block_type in {"bullets", "numbered"}:
                for item_index, item in enumerate(block.get("items") or [], 1):
                    marker = "\u2022" if block_type == "bullets" else f"{item_index}."
                    elements.append(Paragraph(
                        _markdown_inline(str(item)),
                        styles["bullet"],
                        bulletText=marker,
                    ))
                first_narrative = False
            elif block_type == "table":
                headers = [str(cell) for cell in block.get("headers") or []]
                rows = [
                    [str(cell) for cell in row]
                    for row in block.get("rows") or []
                    if isinstance(row, list)
                ]
                if headers and rows:
                    elements.append(_professional_table(
                        headers,
                        rows,
                        table_width=table_width,
                        body_style=styles["table_body"],
                        header_style=styles["table_header"],
                    ))
                    elements.append(Spacer(1, 3 * mm))
                first_narrative = False
        if section_index < len(report_document.get("sections") or []) - 1:
            elements.append(Spacer(1, 4 * mm))
    append_chart_fallback()
def _build_professional_pdf(
    title: str,
    summary: str,
    sections: dict[str, Any],
    *,
    report_document: dict[str, Any] | None = None,
    content_md: str = "",
) -> bytes:
    """Generate a designed PDF from a semantic document, never Markdown flowables."""
    from datetime import datetime
    from backend.services.report_document import build_report_document, normalize_report_document
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer
    _check_reportlab()
    _register_cjk_fonts()
    title_font = _CJK_BOLD_FONT_NAME if _CJK_BOLD_FONT_NAME in pdfmetrics.getRegisteredFontNames() else _CJK_FONT_NAME
    body_font = _CJK_FONT_NAME if _CJK_FONT_NAME in pdfmetrics.getRegisteredFontNames() else "Helvetica"
    if report_document:
        document = normalize_report_document(report_document, title=title, summary=summary)
    else:
        document = build_report_document(content_md, title=title, summary=summary)
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=20 * mm,
        bottomMargin=18 * mm,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        title=title,
        author="AI E-commerce Multi-Agent System",
    )
    base = getSampleStyleSheet()
    styles: dict[str, Any] = {
        "cover_label": ParagraphStyle(
            "ProfessionalCoverLabel", parent=base["Normal"], fontName=title_font,
            fontSize=10, leading=14, textColor=colors.HexColor("#b9d6ef"),
            alignment=TA_LEFT, spaceAfter=8,
        ),
        "cover_title": ParagraphStyle(
            "ProfessionalCoverTitle", parent=base["Title"], fontName=title_font,
            fontSize=24, leading=34, textColor=colors.white,
            alignment=TA_LEFT, spaceAfter=14,
        ),
        "cover_summary_label": ParagraphStyle(
            "ProfessionalCoverSummaryLabel", parent=base["Normal"], fontName=title_font,
            fontSize=11, leading=15, textColor=colors.HexColor("#ffffff"), spaceAfter=5,
        ),
        "cover_summary": ParagraphStyle(
            "ProfessionalCoverSummary", parent=base["Normal"], fontName=body_font,
            fontSize=10.5, leading=18, textColor=colors.HexColor("#e8f1f8"),
            spaceAfter=10,
        ),
        "cover_meta": ParagraphStyle(
            "ProfessionalCoverMeta", parent=base["Normal"], fontName=body_font,
            fontSize=8.5, leading=13, textColor=colors.HexColor("#cbd5e1"),
        ),
        "section": ParagraphStyle(
            "ProfessionalSection", parent=base["Heading1"], fontName=title_font,
            fontSize=15, leading=21, textColor=colors.HexColor("#0f3460"),
            backColor=colors.HexColor("#eaf2f8"), borderColor=colors.HexColor("#8ab6d6"),
            borderWidth=0.6, borderPadding=(7, 8, 7, 8),
            spaceBefore=8, spaceAfter=9,
        ),
        "subheading": ParagraphStyle(
            "ProfessionalSubheading", parent=base["Heading2"], fontName=title_font,
            fontSize=11.5, leading=17, textColor=colors.HexColor("#1a1a2e"),
            spaceBefore=8, spaceAfter=5, keepWithNext=True,
        ),
        "lead": ParagraphStyle(
            "ProfessionalLead", parent=base["Normal"], fontName=body_font,
            fontSize=10, leading=16, textColor=colors.HexColor("#1e293b"),
            backColor=colors.HexColor("#f5f9fc"), borderColor=colors.HexColor("#0f3460"),
            borderWidth=0.7, borderPadding=(7, 9, 7, 9),
            spaceAfter=8,
        ),
        "body": ParagraphStyle(
            "ProfessionalBody", parent=base["Normal"], fontName=body_font,
            fontSize=9.5, leading=15.5, textColor=colors.HexColor("#334155"),
            spaceAfter=6, wordWrap="CJK",
        ),
        "callout": ParagraphStyle(
            "ProfessionalCallout", parent=base["Normal"], fontName=body_font,
            fontSize=9.5, leading=15, textColor=colors.HexColor("#0f3460"),
            backColor=colors.HexColor("#eef6fb"), borderColor=colors.HexColor("#8ab6d6"),
            borderWidth=0.5, borderPadding=(6, 8, 6, 8), spaceAfter=7,
        ),
        "bullet": ParagraphStyle(
            "ProfessionalBullet", parent=base["Normal"], fontName=body_font,
            fontSize=9.5, leading=15, textColor=colors.HexColor("#334155"),
            leftIndent=15, firstLineIndent=-9, bulletIndent=3, spaceAfter=3,
        ),
        "table_header": ParagraphStyle(
            "ProfessionalTableHeader", parent=base["Normal"], fontName=title_font,
            fontSize=7.5, leading=10, textColor=colors.white, spaceAfter=0,
        ),
        "table_body": ParagraphStyle(
            "ProfessionalTableBody", parent=base["Normal"], fontName=body_font,
            fontSize=7.3, leading=10, textColor=colors.HexColor("#334155"), spaceAfter=0,
            wordWrap="CJK",
        ),
        "chart_title": ParagraphStyle(
            "ProfessionalChartTitle", parent=base["Normal"], fontName=title_font,
            fontSize=10.5, leading=15, textColor=colors.HexColor("#0f3460"),
            spaceBefore=5, spaceAfter=4,
        ),
    }
    flat_sections = dict(sections or {})
    nested = flat_sections.pop("_agent_sections", None)
    if isinstance(nested, dict):
        for key, value in nested.items():
            flat_sections.setdefault(key, value)
    chart_insertions: dict[str, list[Any]] = {}
    for agent_key in _AGENT_ORDER:
        data = flat_sections.get(agent_key)
        if not isinstance(data, dict):
            continue
        chart_elements: list[Any] = []
        _render_pdf_charts(chart_elements, agent_key, data, styles["chart_title"])
        if chart_elements:
            chart_insertions[agent_key] = chart_elements
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
    elements: list[Any] = [
        Spacer(1, 8 * mm),
        Paragraph("\u4f01\u4e1a\u7ea7\u7535\u5546\u667a\u80fd\u51b3\u7b56\u62a5\u544a", styles["cover_label"]),
        Paragraph(_markdown_inline(document.get("title") or title), styles["cover_title"]),
        Spacer(1, 5 * mm),
        Paragraph("\u6267\u884c\u6458\u8981", styles["cover_summary_label"]),
        Paragraph(_markdown_inline(document.get("summary") or summary or "\u672c\u62a5\u544a\u5df2\u5b8c\u6210\u3002"), styles["cover_summary"]),
        Spacer(1, 8 * mm),
        Paragraph(
            f"REPORT / {generated_at}<br/>AI E-COMMERCE MULTI-AGENT SYSTEM",
            styles["cover_meta"],
        ),
        PageBreak(),
    ]
    _render_semantic_report(
        elements,
        document,
        styles=styles,
        table_width=doc.width,
        chart_insertions=chart_insertions,
    )
    page_width, page_height = A4
    def draw_cover(canvas, document_template) -> None:
        canvas.saveState()
        canvas.setFillColor(colors.HexColor("#0b1f33"))
        canvas.rect(0, page_height - 142 * mm, page_width, 142 * mm, stroke=0, fill=1)
        canvas.setFillColor(colors.HexColor("#e94560"))
        canvas.rect(0, page_height - 145 * mm, page_width, 3 * mm, stroke=0, fill=1)
        canvas.setFillColor(colors.HexColor("#f4f7fb"))
        canvas.rect(0, 0, page_width, page_height - 145 * mm, stroke=0, fill=1)
        canvas.setStrokeColor(colors.HexColor("#cbd5e1"))
        canvas.line(18 * mm, 15 * mm, page_width - 18 * mm, 15 * mm)
        canvas.setFont(body_font, 7.5)
        canvas.setFillColor(colors.HexColor("#64748b"))
        canvas.drawString(18 * mm, 10 * mm, "CONFIDENTIAL / INTERNAL USE")
        canvas.drawRightString(page_width - 18 * mm, 10 * mm, generated_at)
        canvas.restoreState()
    def draw_body_page(canvas, document_template) -> None:
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#dbe5ee"))
        canvas.line(18 * mm, page_height - 13 * mm, page_width - 18 * mm, page_height - 13 * mm)
        canvas.setFont(body_font, 7.5)
        canvas.setFillColor(colors.HexColor("#64748b"))
        canvas.drawString(18 * mm, page_height - 9 * mm, "AI E-COMMERCE INTELLIGENCE")
        canvas.drawRightString(page_width - 18 * mm, page_height - 9 * mm, str(title)[:42])
        canvas.line(18 * mm, 13 * mm, page_width - 18 * mm, 13 * mm)
        canvas.drawString(18 * mm, 8 * mm, "DATA-DRIVEN DECISION REPORT")
        canvas.drawRightString(page_width - 18 * mm, 8 * mm, f"PAGE {document_template.page}")
        canvas.restoreState()
    doc.build(elements, onFirstPage=draw_cover, onLaterPages=draw_body_page)
    return buffer.getvalue()
class ReportExportService:
    """Generate PDF and DOCX reports from structured sections data."""
    def to_pdf(
        self,
        title: str,
        summary: str,
        sections: dict[str, Any],
        content_md: str = "",
        report_document: dict[str, Any] | None = None,
    ) -> bytes:
        """Render a full PDF report and return the bytes."""
        return _build_professional_pdf(
            title,
            summary,
            sections,
            report_document=report_document,
            content_md=content_md,
        )
    def to_docx(
        self,
        title: str,
        summary: str,
        sections: dict[str, Any],
        content_md: str = "",
        report_document: dict[str, Any] | None = None,
    ) -> bytes:
        """Render a full DOCX report and return the bytes."""
        return _build_docx(
            title,
            summary,
            sections,
            report_document=report_document,
            content_md=content_md,
        )
# ---------------------------------------------------------------------------

#  DOCX semantic helpers (mirror _render_semantic_report for PDF)

# ---------------------------------------------------------------------------

def _collect_docx_charts(agent_key: str, data: dict[str, Any]) -> list[tuple[str, Any]]:
    """Collect (chart_title, PIL.Image) pairs for one agent.
    Mirrors `_render_pdf_charts` but returns data instead of mutating a
    flowable list, so the semantic DOCX renderer can place each chart near
    the heading it belongs to (same logic as the PDF renderer).
    """
    PRODUCT_PRIMARY = "#0f3460"
    ACCENT = "#e94560"
    SUCCESS = "#22c55e"
    WARNING = "#f59e0b"
    GREY = "#94a3b8"
    out: list[tuple[str, Any]] = []
    if agent_key == "product_analysis":
        products = data.get("selected_products", []) or []
        if len(products) >= 2:
            items = products[:10]
            labels = [p.get("title", "") for p in items]
            values = [
                float(p.get("final_score", 0) or p.get("composite_score", 0) or 0)
                for p in items
            ]
            if any(v > 0 for v in values):
                img = _make_bar_chart(
                    "综合评分（越高越好）", labels, values,
                    color_hex=PRODUCT_PRIMARY, value_format="{:.1f}",
                )
                if img is not None:
                    out.append(("Top 10 商品综合评分对比", img))
            cat_dist = (data.get("statistics", {}) or {}).get("category_distribution", {}) or {}
            if len(cat_dist) >= 2:
                img = _make_pie_chart(
                    "类目商品数", list(cat_dist.keys()), list(cat_dist.values())
                )
                if img is not None:
                    out.append(("类目商品分布", img))
    elif agent_key == "trend_forecast":
        forecasts = data.get("product_forecasts", []) or []
        if len(forecasts) >= 2:
            items = forecasts[:10]
            labels = [fc.get("title", "") for fc in items]
            def _sum30(fc):
                f30 = fc.get("forecast_30d") or fc.get("predicted_30d") or []
                if isinstance(f30, list):
                    return float(sum(f30))
                return float(f30 or 0)
            values_30 = [_sum30(fc) for fc in items]
            values_growth = [
                float(fc.get("trend_rate", 0) or fc.get("growth_rate", 0) or 0) * 100
                for fc in items
            ]
            if any(v > 0 for v in values_30):
                img = _make_bar_chart(
                    "30 天预测销量（件）", labels, values_30,
                    color_hex=ACCENT, value_format="{:.0f}",
                )
                if img is not None:
                    out.append(("30 天销量预测对比", img))
            if any(v != 0 for v in values_growth):
                color = SUCCESS if any(v >= 0 for v in values_growth) else ACCENT
                img = _make_bar_chart(
                    "增长率（%）", labels, values_growth,
                    color_hex=color, value_format="{:.1f}",
                )
                if img is not None:
                    out.append(("增长率对比", img))
    elif agent_key == "pricing":
        results = data.get("pricing_results", []) or []
        if len(results) >= 2:
            items = results[:8]
            labels = [r.get("title", "") for r in items]
            cur = [float(r.get("current_price", 0) or 0) for r in items]
            sug = [float(r.get("suggested_price", 0) or 0) for r in items]
            img = _make_grouped_bar_chart(
                "定价对比（USD）", labels, cur, sug,
                legend_a="当前价", legend_b="建议价",
                color_a_hex=GREY, color_b_hex=ACCENT,
                value_format="",
            )
            if img is not None:
                out.append(("定价对比：当前价 vs 建议价", img))
    elif agent_key == "inventory":
        plans = data.get("replenishment_plans", []) or []
        if len(plans) >= 2:
            items = plans[:8]
            labels = [pl.get("title", "") for pl in items]
            values = [float(pl.get("suggested_reorder_qty", 0) or 0) for pl in items]
            if any(v > 0 for v in values):
                img = _make_bar_chart(
                    "建议补货量（件）", labels, values,
                    color_hex=WARNING, value_format="{:.0f}",
                )
                if img is not None:
                    out.append(("建议补货数量", img))
    return out
def _setup_docx_page(doc: Any) -> None:
    """Configure page margins and base font for a professional DOCX report."""
    from docx.shared import Cm, Pt
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
def _render_docx_cover(
    doc: Any,
    document: dict[str, Any],
    title: str,
    summary: str,
) -> None:
    """Render the professional cover page (mirrors _build_professional_pdf)."""
    from datetime import datetime
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
    # Eyebrow label
    eyebrow = doc.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = eyebrow.add_run("企业级电子商智决策报告")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xB9, 0xD6, 0xEF)
    run.bold = True
    # Title (large)
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.add_run(document.get("title") or title or "分析报告")
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(0x0B, 0x1F, 0x33)
    doc.add_paragraph()
    # Summary block
    label = doc.add_paragraph()
    run = label.add_run("执行摘要")
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(0x0B, 0x1F, 0x33)
    body = doc.add_paragraph()
    run = body.add_run(
        document.get("summary") or summary or "本报告已完成。"
    )
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    doc.add_paragraph()
    meta = doc.add_paragraph()
    run = meta.add_run(
        f"REPORT / {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        "AI E-COMMERCE MULTI-AGENT SYSTEM"
    )
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
def _add_docx_table(
    doc: Any,
    headers: list[str],
    rows: list[list[str]],
) -> None:
    """Insert a styled table with bold header row."""
    if not headers or not rows:
        return
    from docx.shared import Pt

    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Light Grid Accent 1"
    except Exception:
        table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, txt in enumerate(headers):
        hdr[i].text = txt
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if i < len(cells):
                cells[i].text = val
def _render_semantic_docx(
    doc: Any,
    document: dict[str, Any],
    *,
    chart_insertions: dict[str, list[tuple[str, Any]]],
) -> None:
    """Walk semantic blocks and render into DOCX with chart insertion.
    Mirrors `_render_semantic_report` for PDF: same heading ordering,
    block type coverage, and heading-matched chart placement.
    """
    import io as _io
    from docx.shared import Inches, Pt, RGBColor
    def append_matching_charts(heading_text: str) -> None:
        matched = [
            agent_key
            for agent_key in list(chart_insertions)
            if _chart_heading_matches(agent_key, heading_text)
        ]
        for agent_key in matched:
            for chart_title, pil_img in chart_insertions.pop(agent_key):
                if pil_img is None:
                    continue
                if chart_title:
                    cap = doc.add_paragraph()
                    run = cap.add_run(chart_title)
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
                buf = _io.BytesIO()
                pil_img.save(buf, format="PNG")
                buf.seek(0)
                doc.add_picture(buf, width=Inches(5.5))
                doc.add_paragraph()
    def append_chart_fallback() -> None:
        if not chart_insertions:
            return
        doc.add_heading("数据可视化", level=1)
        for agent_key, charts in list(chart_insertions.items()):
            label = _SECTION_LABELS.get(agent_key, agent_key)
            doc.add_heading(label, level=2)
            for chart_title, pil_img in charts:
                if pil_img is None:
                    continue
                if chart_title:
                    cap = doc.add_paragraph()
                    run = cap.add_run(chart_title)
                    run.bold = True
                    run.font.size = Pt(10)
                buf = _io.BytesIO()
                pil_img.save(buf, format="PNG")
                buf.seek(0)
                doc.add_picture(buf, width=Inches(5.5))
                doc.add_paragraph()
            chart_insertions.pop(agent_key, None)
    sections = document.get("sections") or []
    for section in sections:
        if not isinstance(section, dict):
            continue
        section_title = str(section.get("title") or "分析结果").strip()
        doc.add_heading(section_title, level=1)
        append_matching_charts(section_title)
        first_narrative = True
        for block in section.get("blocks") or []:
            if not isinstance(block, dict):
                continue
            block_type = str(block.get("type") or "paragraph")
            if block_type == "heading":
                heading_text = str(block.get("text") or "").strip()
                if not heading_text:
                    continue
                doc.add_heading(heading_text, level=2)
                append_matching_charts(heading_text)
                first_narrative = False
            elif block_type == "paragraph":
                text = str(block.get("text") or "").strip()
                if not text:
                    continue
                if first_narrative:
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    run.bold = True
                    first_narrative = False
                else:
                    doc.add_paragraph(text)
            elif block_type == "callout":
                text = str(block.get("text") or "").strip()
                if text:
                    doc.add_paragraph(text)
            elif block_type in {"bullets", "numbered"}:
                style = "List Bullet" if block_type == "bullets" else "List Number"
                for item in block.get("items") or []:
                    doc.add_paragraph(str(item), style=style)
                first_narrative = False
            elif block_type == "table":
                headers = [str(c) for c in block.get("headers") or []]
                rows = [
                    [str(c) for c in row]
                    for row in block.get("rows") or []
                    if isinstance(row, list)
                ]
                _add_docx_table(doc, headers, rows)
                if headers and rows:
                    doc.add_paragraph()
    append_chart_fallback()

